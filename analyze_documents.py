#!/usr/bin/env python3
"""
Script to analyze AnthosKS project documents
"""
import openpyxl
from docx import Document
import json

# Analyze Excel file
print("=" * 80)
print("ANALYZING EXCEL TRACKER")
print("=" * 80)

wb = openpyxl.load_workbook('/Users/johanstromquist/Dropbox/Coding/AnthosKS/references/Anthos_ABRS_Global Deployment Tracker_METRICS_20250826.xlsx')
print(f"\nWorksheets: {wb.sheetnames}")

for sheet_name in wb.sheetnames:
    ws = wb[sheet_name]
    print(f"\n--- Sheet: {sheet_name} ---")
    print(f"Dimensions: {ws.dimensions}")
    print(f"Max row: {ws.max_row}, Max column: {ws.max_column}")

    # Print first 5 rows
    print("\nFirst 5 rows:")
    for i, row in enumerate(ws.iter_rows(max_row=5, values_only=True), 1):
        print(f"Row {i}: {row}")

    if ws.max_row > 5:
        print(f"\n... ({ws.max_row - 5} more rows)")

wb.close()

# Analyze Word document
print("\n" + "=" * 80)
print("ANALYZING WORD DOCUMENT")
print("=" * 80)

doc = Document('/Users/johanstromquist/Dropbox/Coding/AnthosKS/references/Anthos_MOV Rpt Summary.docx')

print(f"\nTotal paragraphs: {len(doc.paragraphs)}")
print(f"Total tables: {len(doc.tables)}")

print("\nFirst 10 paragraphs:")
for i, para in enumerate(doc.paragraphs[:10], 1):
    if para.text.strip():
        print(f"{i}. {para.text[:100]}...")

if doc.tables:
    print(f"\n--- Table Analysis ---")
    for i, table in enumerate(doc.tables, 1):
        print(f"\nTable {i}: {len(table.rows)} rows x {len(table.columns)} columns")
        print("First 3 rows:")
        for j, row in enumerate(table.rows[:3], 1):
            cells = [cell.text.strip()[:30] for cell in row.cells]
            print(f"  Row {j}: {cells}")
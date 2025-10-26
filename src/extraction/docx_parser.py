"""DOCX text extraction with layout preservation."""

import docx
from pathlib import Path
from typing import Dict
import logging

logger = logging.getLogger(__name__)


class DOCXParser:
    """Extract text from DOCX files with layout preservation."""

    def __init__(self, preserve_tables: bool = True):
        self.preserve_tables = preserve_tables

    def extract_text(self, docx_path: Path) -> str:
        """
        Extract all text from DOCX with section markers.

        Args:
            docx_path: Path to DOCX file

        Returns:
            Full text with section separators
        """
        logger.info(f"Extracting text from: {docx_path}")

        try:
            doc = docx.Document(docx_path)
            text_parts = []
            page_count = 0  # DOCX doesn't have explicit page concept, we'll estimate

            for i, para in enumerate(doc.paragraphs):
                text = para.text.strip()

                if text:
                    # Add page markers every ~50 paragraphs (rough estimate)
                    if i % 50 == 0:
                        page_count += 1
                        text_parts.append(f"--- PAGE {page_count} ---")

                    text_parts.append(text)

            # Extract tables if requested
            if self.preserve_tables:
                for table in doc.tables:
                    table_text = self._extract_table_text(table)
                    if table_text:
                        text_parts.append("\n--- TABLE ---")
                        text_parts.append(table_text)

            full_text = "\n".join(text_parts)

            logger.info(
                f"Extracted {len(doc.paragraphs)} paragraphs, "
                f"{len(doc.tables)} tables, "
                f"{len(full_text)} characters"
            )

            return full_text

        except Exception as e:
            logger.error(f"DOCX extraction failed: {e}")
            raise

    def _extract_table_text(self, table) -> str:
        """Extract text from a table."""
        table_rows = []

        for row in table.rows:
            row_text = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    row_text.append(cell_text)

            if row_text:
                table_rows.append(" | ".join(row_text))

        return "\n".join(table_rows)

    def extract_metadata(self, docx_path: Path) -> Dict:
        """Extract DOCX metadata."""
        doc = docx.Document(docx_path)

        return {
            "paragraph_count": len(doc.paragraphs),
            "table_count": len(doc.tables),
            "file_size_mb": docx_path.stat().st_size / (1024 * 1024)
        }
